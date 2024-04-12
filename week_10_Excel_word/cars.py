import docx

document = docx.Document()  # empty document we are adding to

electric_cars = {
    'Chevrolet': 'Volt',
    'Nissan': 'Leaf',
    'Tesla': 'Model S',
    'Volkswagen': 'e-Golf'
}  # dictionary we are using to build the word file

document.add_paragraph('Electric Cars', 'Heading 1')
for make, model in electric_cars.items():
    document.add_paragraph(make, 'Heading 3')  # adds heading with variable indicated
    document.add_paragraph(f'An electric car by {make} is {model} ')  # format strings to write a neat sentence
    # document.add_paragraph('make: ' + make)
    # document.add_paragraph('Model: ' + model)

document.save('electric_cars.docx')  # saves document for viewing in word
