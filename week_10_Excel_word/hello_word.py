import docx

document = docx.Document()

document.add_paragraph('Hello Word!',
                       'Title')  # must be exact, style type must go after what we are trying to set there
document.add_paragraph('By Taylor', 'Heading 1')  # must be exact
document.add_paragraph('This is a word document created with Python and python-docx')
# saved the Style names slide as a screenshot for future ref
document.add_paragraph('Automate the Boring Stuff.', 'Quote')
document.add_paragraph('This is the book we use for Programming Logic and Design', 'Body Text')

document.add_paragraph('List of Favorite Colors', 'Heading 2')  # watch your comma placements!

favorite_colors = ['Green', 'Pink', 'Black']

for color in favorite_colors:
    document.add_paragraph(color, 'List Bullet')

document.save('hello_word.docx')

# tip: use this code as a skeleton for Lab pt 2~
